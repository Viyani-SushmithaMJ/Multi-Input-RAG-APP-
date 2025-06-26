from langchain_community.embeddings import OpenAIEmbeddings
from langchain_community.vectorstores import FAISS
from langchain_community.chat_models import ChatOpenAI
from langchain.chains import RetrievalQAWithSourcesChain
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from youtube_transcript_api import YouTubeTranscriptApi
import time
import docx
import PyPDF2
import tiktoken
import openai
import whisper
import re
import hashlib
import tempfile
import streamlit as st
import os
import csv
from datetime import datetime

os.environ["KMP_DUPLICATE_LIB_OK"] = "TRUE"

# --- API Key ---
openai_api_key = os.getenv("OPENAI_API_KEY")
openai.api_key = openai_api_key

# --- Constants ---
WHISPER_MODELS = ["tiny", "base", "small", "medium", "large"]
MODEL_COST = {
    "gpt-3.5-turbo": 0.0015 / 1000,
    "gpt-4": 0.03 / 1000,
}
FAISS_DIR = "faiss_index"
os.makedirs(FAISS_DIR, exist_ok=True)

FEEDBACK_CSV = "feedback_log.csv"

# --- Utility Functions ---


def extract_video_id(url):
    match = re.search(r"(?:v=|\/)([0-9A-Za-z_-]{11})", url)
    return match.group(1) if match else None


def get_transcript_from_youtube(url):
    video_id = extract_video_id(url)
    transcript = YouTubeTranscriptApi.get_transcript(video_id)
    return " ".join([entry['text'] for entry in transcript])


def transcribe_video(file_path, model_size):
    model = whisper.load_model(model_size)
    result = model.transcribe(file_path)
    return result["text"]


def parse_srt_or_vtt(file):
    content = file.read().decode("utf-8")
    lines = content.splitlines()
    text_lines = [line for line in lines if not re.match(
        r"^\d+$|^(\d{2}:\d{2}:\d{2},\d{3})", line)]
    return " ".join(text_lines)


def read_txt(file):
    return file.read().decode("utf-8")


def read_pdf(file):
    reader = PyPDF2.PdfReader(file)
    return " ".join(page.extract_text() for page in reader.pages if page.extract_text())


def read_docx(file):
    doc = docx.Document(file)
    return "\n".join([p.text for p in doc.paragraphs])


def prepare_documents(text, source=""):
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000, chunk_overlap=200)
    chunks = splitter.split_text(text)
    return [Document(page_content=chunk, metadata={"source": source}) for chunk in chunks]


def estimate_tokens_and_cost(text, model_name):
    enc = tiktoken.encoding_for_model(model_name)
    tokens = len(enc.encode(text))
    cost = tokens * MODEL_COST.get(model_name, 0.0015 / 1000)
    return tokens, cost


def get_md5(text):
    return hashlib.md5(text.encode("utf-8")).hexdigest()


def build_or_load_retriever(docs, identifier):
    faiss_path = os.path.join("faiss_index", identifier)
    embeddings = OpenAIEmbeddings(openai_api_key=openai_api_key)

    if os.path.exists(faiss_path):
        return FAISS.load_local(
            faiss_path,
            embeddings,
            allow_dangerous_deserialization=True
        ).as_retriever()
    else:
        vectorstore = FAISS.from_documents(docs, embeddings)
        vectorstore.save_local(faiss_path)
        return vectorstore.as_retriever()


def get_qa_chain(retriever, model_name="gpt-3.5-turbo"):
    llm = ChatOpenAI(model_name=model_name, temperature=0.3,
                     openai_api_key=openai_api_key)
    return RetrievalQAWithSourcesChain.from_chain_type(llm=llm, retriever=retriever, chain_type="stuff")


def save_feedback(question, answer, rating, comment=""):
    header = ["timestamp", "question", "answer", "rating", "comment"]
    row = [datetime.now().isoformat(), question, answer, rating, comment]

    file_exists = os.path.isfile(FEEDBACK_CSV)
    with open(FEEDBACK_CSV, mode="a", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        if not file_exists:
            writer.writerow(header)
        writer.writerow(row)


def load_feedback_stats():
    if not os.path.isfile(FEEDBACK_CSV):
        return None
    ratings = {"👍 Yes": 0, "👎 No": 0, "🤷 Not Sure": 0}
    total = 0
    with open(FEEDBACK_CSV, mode="r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rating = row.get("rating")
            if rating in ratings:
                ratings[rating] += 1
                total += 1
    return total, ratings


# ---  UI ---

st.set_page_config(page_title="📚 Multi-Input Q&A", layout="centered")
st.title("🎧 Multi-Modal Q&A with LangChain + Whisper + FAISS Caching")

# Display feedback stats at top
stats = load_feedback_stats()
if stats:
    total, ratings = stats
    st.markdown("#### 🧾 Feedback Summary")
    st.write(f"Total feedback received: **{total}**")
    st.write(f"👍 Accurate: **{ratings['👍 Yes']}**")
    st.write(f"👎 Inaccurate: **{ratings['👎 No']}**")
    st.write(f"🤷 Not Sure: **{ratings['🤷 Not Sure']}**")
    st.markdown("---")

st.markdown("#### 📺 Enter YouTube URLs (comma/newline separated)")
youtube_input = st.text_area("YouTube Links")

st.markdown("#### 📂 Upload documents, videos, or audio files")
uploaded_files = st.file_uploader(
    "Upload .mp4/.mov/.mp3/.wav/.m4a/.srt/.vtt/.txt/.pdf/.docx files",
    accept_multiple_files=True
)

whisper_model_choice = st.selectbox(
    "🎙️ Whisper Model", WHISPER_MODELS, index=1)
model_choice = st.selectbox("🤖 LLM", ["gpt-3.5-turbo", "gpt-4"])
user_question = st.text_input("❓ Ask your question")

if st.button("Answer"):
    if not user_question:
        st.error("Please enter a question.")
    elif not openai_api_key:
        st.error("Missing OpenAI API key in environment variable `OPENAI_API_KEY`.")
    else:
        with st.spinner("⏳ Processing..."):
            all_docs = []
            audio_video_exts = (".mp4", ".mov", ".mp3", ".wav", ".m4a")

            if youtube_input:
                urls = [u.strip() for u in re.split(
                    r"[\n,]+", youtube_input) if u.strip()]
                for url in urls:
                    try:
                        text = get_transcript_from_youtube(url)
                        all_docs.extend(prepare_documents(
                            text, source=f"YouTube: {url}"))
                    except Exception as e:
                        st.warning(
                            f"❌ Failed to fetch transcript from: {url} — {e}")

            for file in uploaded_files:
                filename = file.name.lower()
                try:
                    if filename.endswith(".txt"):
                        text = read_txt(file)
                    elif filename.endswith((".srt", ".vtt")):
                        text = parse_srt_or_vtt(file)
                    elif filename.endswith(".pdf"):
                        text = read_pdf(file)
                    elif filename.endswith(".docx"):
                        text = read_docx(file)
                    elif filename.endswith(audio_video_exts):
                        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(filename)[1]) as temp_file:
                            temp_file.write(file.read())
                            temp_path = temp_file.name
                        text = transcribe_video(
                            temp_path, whisper_model_choice)
                        os.unlink(temp_path)
                    else:
                        st.warning(f"Unsupported file type: {filename}")
                        continue

                    all_docs.extend(prepare_documents(text, source=filename))
                except Exception as e:
                    st.warning(f"❌ Error reading file {filename}: {e}")

            if not all_docs:
                st.error("❌ No valid content found.")
            else:
                full_text = " ".join([doc.page_content for doc in all_docs])
                text_hash = get_md5(full_text)
                retriever = build_or_load_retriever(all_docs, text_hash)
                chain = get_qa_chain(retriever, model_choice)

                tokens, cost = estimate_tokens_and_cost(
                    full_text + user_question, model_choice)
                start_time = time.time()
                result = chain.invoke({"question": user_question})
                elapsed = time.time() - start_time

                # --- Result ---
                st.success("✅ Answer:")
                st.write(result["answer"])

                st.markdown("#### 📌 Sources Used:")
                for src in set(result.get("sources", "").split(",")):
                    if src.strip():
                        st.markdown(f"- {src.strip()}")

                # Show top transcript snippet (sanity check)
                top_docs = retriever.get_relevant_documents(user_question)
                if top_docs:
                    st.markdown("#### 🔍 Transcript Snippet (Sanity Check):")
                    st.code(top_docs[0].page_content[:500] + "...")

                st.markdown("#### 💰 Token Usage Estimate:")
                st.write(f"Tokens used: `{tokens}`")
                st.write(f"Estimated cost: `${cost:.6f}` USD")
                st.write(f"Response time: `{elapsed:.2f}` seconds")

                # --- Feedback Form ---
                st.markdown("#### 🧪 Rate this Answer")
                rating = st.radio("Is the answer accurate?", [
                                  "👍 Yes", "👎 No", "🤷 Not Sure"], key="rating")

                comment = st.text_area(
                    "Additional feedback (optional):", key="comment")

                if st.button("Submit Feedback"):
                    save_feedback(
                        user_question, result["answer"], rating, comment)
                    st.success("Thank you for your feedback!")
